import markdown2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import re

def clean_html_text(html_text):
    """Clean HTML text by removing problematic tags and entities"""
    text = re.sub(r'<[^>]+>', '', html_text)
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&quot;', '"')
    text = text.replace('&#39;', "'")
    return text.strip()

def markdown_to_docx_elements(markdown_text, doc):
    """Convert markdown text to DOCX elements, rendering all content as-is (including TOC)."""
    html_content = markdown2.markdown(markdown_text, extras=['tables', 'fenced-code-blocks', 'code-friendly'])
    lines = html_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('<h1>'):
            text = re.search(r'<h1>(.*?)</h1>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    heading = doc.add_heading(clean_text, level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('<h2>'):
            text = re.search(r'<h2>(.*?)</h2>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    doc.add_heading(clean_text, level=2)
        elif line.startswith('<h3>'):
            text = re.search(r'<h3>(.*?)</h3>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    doc.add_heading(clean_text, level=3)
        elif line.startswith('<h4>'):
            text = re.search(r'<h4>(.*?)</h4>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    doc.add_heading(clean_text, level=4)
        elif line.startswith('<h5>'):
            text = re.search(r'<h5>(.*?)</h5>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    doc.add_heading(clean_text, level=5)
        elif line.startswith('<h6>'):
            text = re.search(r'<h6>(.*?)</h6>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    doc.add_heading(clean_text, level=6)
        elif line.startswith('<p>'):
            text = re.search(r'<p>(.*?)</p>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    paragraph = doc.add_paragraph(clean_text)
                    paragraph.paragraph_format.space_after = Pt(12)
        elif line.startswith('<pre>'):
            code_text = ""
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('</pre>'):
                code_text += lines[i] + '\n'
                i += 1
            if code_text.strip():
                clean_code = clean_html_text(code_text)
                if clean_code:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(clean_code)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.paragraph_format.left_indent = Inches(0.5)
                    code_para.paragraph_format.right_indent = Inches(0.5)
                    code_para.paragraph_format.space_after = Pt(12)
        elif line.startswith('<code>'):
            text = re.search(r'<code>(.*?)</code>', line, re.DOTALL)
            if text:
                clean_text = clean_html_text(text.group(1))
                if clean_text:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(clean_text)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.paragraph_format.space_after = Pt(6)
        elif line.startswith('<blockquote>'):
            quote_text = ""
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('</blockquote>'):
                quote_text += lines[i] + '\n'
                i += 1
            if quote_text.strip():
                clean_quote = clean_html_text(quote_text)
                if clean_quote:
                    quote_para = doc.add_paragraph(clean_quote)
                    quote_para.paragraph_format.left_indent = Inches(0.5)
                    quote_para.paragraph_format.right_indent = Inches(0.5)
                    quote_para.paragraph_format.space_after = Pt(12)
                    for run in quote_para.runs:
                        run.italic = True
        elif line.startswith('<ul>'):
            list_items = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('</ul>'):
                if lines[i].strip().startswith('<li>'):
                    item_text = re.search(r'<li>(.*?)</li>', lines[i].strip(), re.DOTALL)
                    if item_text:
                        clean_item = clean_html_text(item_text.group(1))
                        if clean_item:
                            list_items.append(clean_item)
                i += 1
            if list_items:
                for item in list_items:
                    list_para = doc.add_paragraph(item, style='List Bullet')
                    list_para.paragraph_format.space_after = Pt(6)
        elif line.startswith('<ol>'):
            list_items = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('</ol>'):
                if lines[i].strip().startswith('<li>'):
                    item_text = re.search(r'<li>(.*?)</li>', lines[i].strip(), re.DOTALL)
                    if item_text:
                        clean_item = clean_html_text(item_text.group(1))
                        if clean_item:
                            list_items.append(clean_item)
                i += 1
            if list_items:
                for item in list_items:
                    list_para = doc.add_paragraph(item, style='List Number')
                    list_para.paragraph_format.space_after = Pt(6)
        elif line.startswith('<li>'):
            item_text = re.search(r'<li>(.*?)</li>', line, re.DOTALL)
            if item_text:
                clean_item = clean_html_text(item_text.group(1))
                if clean_item:
                    if re.match(r'^\d+\.', clean_item.strip()):
                        list_para = doc.add_paragraph(clean_item, style='List Number')
                    else:
                        list_para = doc.add_paragraph(clean_item, style='List Bullet')
                    list_para.paragraph_format.space_after = Pt(6)
        elif line and not line.startswith('<') and not line.startswith('</'):
            clean_text = clean_html_text(line)
            if clean_text:
                if re.match(r'^\d+\.', clean_text.strip()):
                    list_para = doc.add_paragraph(clean_text, style='List Number')
                    list_para.paragraph_format.space_after = Pt(6)
                else:
                    paragraph = doc.add_paragraph(clean_text)
                    paragraph.paragraph_format.space_after = Pt(12)
        i += 1

def convert_markdown_to_docx(markdown_text, output_filename=None):
    """
    Convert markdown text to a DOCX file and return the filename.
    If output_filename is not provided, generate one with a timestamp in the output folder inside deep_research.
    """
    if not markdown_text:
        print("Error: No markdown content provided")
        return None
    try:
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"research_report_{timestamp}.docx"
        doc = Document()
        doc.core_properties.title = "Research Report"
        doc.core_properties.author = "Deep Research System"
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        markdown_to_docx_elements(markdown_text, doc)
        output_dir = os.path.join(os.path.dirname(__file__), 'output')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        print(f"DOCX created successfully: {output_path}")
        return output_path
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return None

def test_conversion():
    """Test function to validate the markdown to DOCX conversion"""
    test_markdown = """
# Test Research Report

## Table of Contents
1. **Introduction**  
2. **Key Points**  
3. **Code Example**  
    1. **Major Technology Companies**  
    2. **Emerging Startups**  
    3. **AI Frameworks**  
4. **Conclusion**  

## Introduction
This is a test report to validate the markdown to DOCX conversion functionality.

### Key Points
- Point 1: Testing basic functionality
- Point 2: Testing code blocks
- Point 3: Testing formatting

## Code Example
Here's some example code:

```python
def hello_world():
    print("Hello, World!")
    return "Success"
```

## Conclusion
This test validates that the DOCX conversion is working correctly.
"""
    
    print("Testing markdown to DOCX conversion...")
    result = convert_markdown_to_docx(test_markdown, "test_report.docx")
    
    if result:
        print("✅ Test passed! DOCX created successfully.")
        return True
    else:
        print("❌ Test failed! DOCX creation failed.")
        return False

if __name__ == "__main__":
    # Run test when script is executed directly
    test_conversion()
    # pass 